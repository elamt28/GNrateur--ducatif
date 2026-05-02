import streamlit as st
import traceback

# --- 1. CONFIGURATION DE LA PAGE (SÉCURISÉE) ---
try:
    st.set_page_config(page_title="EduForge Pro V36", layout="wide", page_icon="⚡")
except Exception:
    pass # Tolérance aux rechargements internes

# --- 2. BOUCLIER GLOBAL ANTI-CRASH ---
try:
    import google.generativeai as genai
    from io import BytesIO
    
    # --- VÉRIFICATION DU MODULE WORD ---
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        HAS_DOCX = True
    except ImportError:
        HAS_DOCX = False

    # --- MÉMOIRE DE SESSION ---
    if 'cours_memoire' not in st.session_state:
        st.session_state.cours_memoire = None
    if 'sujet_memoire' not in st.session_state:
        st.session_state.sujet_memoire = None

    # --- MOTEUR DE GÉNÉRATION V36 (PÉDAGOGIE EXPERTE CFA CHARTRES) ---
    def forger_cours_v36(formation, sujet, lieu, moteur, api_key):
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(moteur)
        
        prompt = f"""
        Agis en expert pédagogique du CFA Interpro de Chartres. 
        Module FOAD de 60 minutes minimum. Formation : {formation}. Sujet : {sujet}. 
        Localisation : {lieu} (Chartres/Champhol).
        
        RÈGLES ABSOLUES ET CERTIFIÉES :
        1. NE SALUE PAS au début du cours (Pas de "Bonjour", "Bienvenue", etc.). Rentre dans le vif du sujet.
        2. Ne cite JAMAIS le prénom 'Manu'.
        3. Ton ludique, humour technique et jeux de mots.
        4. Exactitude certifiée : auto-critique la logique de tes arguments, n'invente rien, pas de réponses superficielles.
        5. NE PROPOSE PAS de diapositives ou PowerPoint.
        6. Respecte scrupuleusement la structure du plan avec missions, exercices et synthèses.
        
        STRUCTURE EXHAUSTIVE : 
        - # 🎓 TITRE DU MODULE
        - ## 🎯 RÉFÉRENTIEL MÉTIER
        - ## 🎬 SCÉNARIO PRINCIPAL (Humour local)
        - ## 📖 CŒUR TECHNIQUE CERTIFIÉ (Logique argumentée)
        - ## 🔄 ET SI LA SITUATION CHANGEAIT ? (Transfert de compétence)
        - ## 📝 ÉVALUATION SOMMATIVE (/20)
        - ## 👨‍🏫 CORRECTION & AUTO-CRITIQUE (Justifie tes réponses pour prouver leur exactitude)
        """
        return model.generate_content(prompt).text

    # --- EXPORT WORD V36 ---
    def generer_docx_v36(titre, contenu):
        if not HAS_DOCX: return None
        doc = Document()
        t = doc.add_heading(titre, 0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for ligne in contenu.split('\n'):
            if ligne.startswith('# '):
                doc.add_heading(ligne.replace('# ', ''), level=1)
            elif ligne.startswith('## '):
                p = doc.add_heading(ligne.replace('## ', ''), level=2)
                if "CHANGEAIT" in ligne:
                    p.runs[0].font.color.rgb = RGBColor(0, 153, 153) # Cyan
            else:
                doc.add_paragraph(ligne)
                
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # --- INTERFACE UTILISATEUR ---
    st.title("⚡ EduForge Pro : V36 (Structure Anti-Crash)")
    st.markdown("*L'ingénierie certifiée du CFA Interpro de Chartres*")

    # Zone de notification statique (Empêche l'erreur removeChild)
    zone_messages = st.empty()

    with st.sidebar:
        st.header("🔑 Configuration")
        api_key = st.text_input("Clé Gemini :", type="password")
        if api_key:
            moteur = st.selectbox("Moteur IA :", ["gemini-1.5-flash", "gemini-2.5-flash"])
        
        st.divider()
        
        # CORRECTIF MAJEUR : Utilisation d'un formulaire pour bloquer les rechargements intempestifs
        with st.form("formulaire_forge"):
            formations = [
                "BTS Maintenance Véhicule", 
                "Bac Pro Maintenance Véhicule (2de, 1re, Term)", 
                "Carrossier/Peintre",
                "BM Boulanger",
                "BP Boulanger", 
                "BP Boucher", 
                "CAP EPC", 
                "BP Coiffure", 
                "AMLHR"
            ]
            f_sel = st.selectbox("Formation visée :", formations)
            sujet_in = st.text_input("Thème technique :")
            lieu_in = st.text_input("Lieu du scénario :", value="Chartres / Champhol")
            
            bouton_lancer = st.form_submit_button("🚀 Forger le Module")

    # --- LOGIQUE DE DÉCLENCHEMENT ---
    if bouton_lancer:
        if api_key and sujet_in:
            # Le spinner est maintenant dans la zone principale, sécurisé.
            with zone_messages:
                with st.spinner("Analyse des compétences et vérification d'exactitude..."):
                    res = forger_cours_v36(f_sel, sujet_in, lieu_in, moteur, api_key)
                    st.session_state.cours_memoire = res
                    st.session_state.sujet_memoire = sujet_in
        else:
            zone_messages.error("⚠️ Veuillez renseigner votre Clé API et un Sujet.")

    # --- ZONE D'AFFICHAGE SÉCURISÉE ---
    zone_contenu = st.container()
    
    with zone_contenu:
        if st.session_state.cours_memoire:
            st.success(f"✅ Module '{st.session_state.sujet_memoire}' forgé avec exactitude.")
            
            if HAS_DOCX:
                data = generer_docx_v36(st.session_state.sujet_memoire, st.session_state.cours_memoire)
                # Le bouton de téléchargement a maintenant l'option use_container_width pour stabiliser le DOM
                st.download_button(
                    label="📥 Télécharger le Document WORD", 
                    data=data, 
                    file_name=f"Cours_FOAD_{st.session_state.sujet_memoire}.docx", 
                    key="dl_docx_v36_final",
                    use_container_width=True
                )
            else:
                st.warning("⚠️ Module Word désactivé (Erreur d'import).")
                
            st.divider()
            st.markdown(st.session_state.cours_memoire)

# --- FIN DU BOUCLIER ---
except BaseException as e:
    st.error("🚨 DÉFAILLANCE CRITIQUE INTERCEPTÉE PAR LA BOÎTE NOIRE V36")
    st.warning("L'application a évité le crash complet. Voici le rapport technique :")
    st.code(traceback.format_exc())
